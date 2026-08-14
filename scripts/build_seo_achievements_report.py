from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/seo-achievements-report-arabic.docx")


def set_document_rtl(doc: Document) -> None:
    settings = doc.settings._element
    bidi = settings.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        settings.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_run_font(run, size=11, bold=False, color="22262B") -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_fonts.set(qn("w:cs"), "Arial")
    r_fonts.set(qn("w:eastAsia"), "Arial")
    rtl = r_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        r_pr.append(rtl)
    rtl.set(qn("w:val"), "1")


def add_paragraph(doc, text, size=11, bold=False, color="22262B", after=6):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    size = 18 if level == 1 else 14
    color = "3D5258" if level == 1 else "526970"
    p = add_paragraph(doc, text, size=size, bold=True, color=color, after=4)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("• " + text)
    set_run_font(run, size=10.5)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_text(cell, text, bold=False, color="22262B", fill=None):
    cell.text = ""
    if fill:
        shade_cell(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    set_rtl(p)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.2, bold=bold, color=color)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    for index, header in enumerate(headers):
        set_cell_text(table.cell(0, index), header, bold=True, color="FFFFFF", fill="526970")
    for row_data in rows:
        row = table.add_row().cells
        for index, value in enumerate(row_data):
            set_cell_text(row[index], value, bold=index == 0, color="3D5258" if index == 0 else "22262B")
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    set_document_rtl(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:cs"), "Arial")

    add_paragraph(doc, "تقرير إنجازات تحسين محركات البحث SEO", size=24, bold=True, color="3D5258", after=2)
    add_paragraph(doc, "منصة الديوان للاستشارات الهندسية", size=15, bold=True, color="C89B3C", after=10)
    add_paragraph(
        doc,
        "هذا التقرير يعرض مميزات SEO المدعومة في المنصة بعد التحديثات الأخيرة، بصياغة إنجازات واضحة يمكن مشاركتها مع الإدارة أو العميل.",
        size=11.5,
        color="5B6067",
        after=12,
    )

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    mark_header_row(callout.rows[0])
    set_cell_text(
        callout.cell(0, 0),
        "الخلاصة: المنصة أصبحت تدعم طبقة SEO مركزية تشمل العناوين، الأوصاف، Canonical، Robots، Open Graph، Twitter Cards، Structured Data، وخريطة موقع ديناميكية.",
        bold=True,
        color="3D5258",
        fill="F7F4EC",
    )
    doc.add_paragraph()

    add_heading(doc, "نظرة عامة على الإنجاز", 1)
    bullets = [
        "بناء مكوّن SEO موحد يخدم كل الصفحات العامة بدل تكرار الأكواد داخل كل صفحة.",
        "تفعيل وسوم البحث الأساسية: العنوان، الوصف، canonical، robots، ووسوم المشاركة الاجتماعية.",
        "إضافة بيانات منظمة JSON-LD لأنواع المحتوى الأساسية: المؤسسة، الموقع، المقالات، الخدمات، المشاريع، الصور، الفيديوهات، والأسئلة الشائعة.",
        "إضافة sitemap.xml ديناميكي يجمع الصفحات والمحتوى المنشور فقط.",
        "تحسين سياسة الفهرسة حتى لا تظهر صفحات الدخول والصيانة في نتائج البحث.",
        "تحسين الصور من ناحية النص البديل، lazy loading، وصورة الهيرو ذات الأولوية العالية.",
    ]
    for item in bullets:
        add_bullet(doc, item)

    add_heading(doc, "مميزات SEO المدعومة بالمنصة", 1)
    features = [
        ("SEO Titles", "مدعومة", "العنوان يدمج اسم الموقع تلقائيًا ويتعامل مع عناوين الصفحات والمحتوى."),
        ("Meta Description", "مدعومة", "وصف لكل صفحة مع fallback ذكي في حال لم يتم إدخال وصف مخصص."),
        ("Canonical URLs", "مدعومة", "رابط أساسي لكل صفحة مع دعم الحقول المخزنة للمقالات والمشاريع والصور والفيديوهات."),
        ("Robots Meta", "مدعومة", "تحكم بالفهرسة، مع noindex لصفحات البحث والفلاتر والدخول والصيانة."),
        ("Open Graph", "مدعومة", "تحسين شكل الرابط عند المشاركة في واتساب وفيسبوك ولينكدإن وغيرها."),
        ("Twitter Cards", "مدعومة", "بطاقات مشاركة كبيرة للصورة والعنوان والوصف."),
        ("Structured Data", "مدعومة", "JSON-LD لأنواع المحتوى الرئيسية لمساعدة محركات البحث على فهم الصفحة."),
        ("Sitemap ديناميكي", "مدعومة", "خريطة موقع تتحدث من قاعدة البيانات وتعرض المحتوى العام المنشور فقط."),
        ("Robots.txt", "مدعومة", "يسمح بالزحف ويعلن رابط خريطة الموقع."),
        ("Image SEO", "مدعومة", "نصوص بديلة وصور محسنة وتحميل كسول للصور غير الحرجة."),
        ("Internal Links", "مدعومة", "تحسين الروابط الداخلية وإصلاح رابط معرض الصور المكسور."),
        ("Dynamic Content SEO", "مدعومة", "المقالات والخدمات والمشاريع والصور والفيديوهات تستفيد من حقول SEO في لوحة التحكم."),
    ]
    add_table(doc, ["الميزة", "الحالة", "الأثر المتوقع"], features)

    add_heading(doc, "تغطية الصفحات والمحتوى", 1)
    coverage = [
        ("الصفحة الرئيسية", "SEO أساسي + Organization + WebSite + FAQPage عند توفر الأسئلة."),
        ("الخدمات", "وصف فهرس الخدمات + Service Schema في صفحة كل خدمة."),
        ("المشاريع", "وصف الفهرس + CreativeWork Schema + canonical مخصص عند توفره."),
        ("المقالات", "Article Schema + بيانات النشر والتصنيف والكلمات المفتاحية."),
        ("الصور", "ImageObject Schema + احترام خيار indexable + alt text."),
        ("الفيديوهات", "VideoObject Schema + صورة مصغرة + مدة + تاريخ نشر عند توفره."),
        ("الصفحات القانونية", "عناوين وأوصاف وروابط canonical مناسبة."),
        ("الدخول والصيانة", "noindex لمنع ظهور صفحات غير تسويقية في البحث."),
    ]
    add_table(doc, ["النطاق", "الدعم المطبق"], coverage)

    add_heading(doc, "كيف يخدم هذا ظهور الموقع؟", 1)
    add_paragraph(doc, "الطبقة المركزية تقلل الأخطاء وتضمن أن أي صفحة جديدة تبدأ من أساس SEO صحيح بدون الحاجة لتكرار الأكواد يدويًا.")
    add_paragraph(doc, "البيانات المنظمة تساعد محركات البحث على فهم نوع المحتوى، وهذا يرفع جاهزية الموقع للنتائج الغنية عندما تتوفر شروط Google.")
    add_paragraph(doc, "خريطة الموقع الديناميكية تساعد على اكتشاف الصفحات والمحتوى المنشور بسرعة، وتستبعد المحتوى غير العام أو غير القابل للفهرسة.")

    add_heading(doc, "ملاحظات تشغيلية مهمة", 1)
    for item in [
        "قبل الإطلاق يجب ضبط APP_URL على الدومين النهائي حتى تكون روابط sitemap وcanonical دقيقة.",
        "كلما تم تعبئة حقول SEO من لوحة التحكم، ستظهر تلقائيًا في الصفحات الديناميكية.",
        "التحديثات لا تغني عن جودة المحتوى نفسه، لكنها توفر الأساس التقني الصحيح للفهرسة والظهور.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "الخلاصة", 1)
    add_paragraph(
        doc,
        "المنصة أصبحت تمتلك أساس SEO تقني متماسك وقابل للتوسع: وسوم موحدة، بيانات منظمة، خريطة موقع ديناميكية، وتحكم أفضل بالفهرسة. هذا يجعل الموقع جاهزًا نظريًا لفهرسة أوضح وظهور أفضل عند اكتمال المحتوى والدومين النهائي.",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    set_rtl(footer, WD_ALIGN_PARAGRAPH.CENTER)
    run = footer.add_run("الديوان للاستشارات الهندسية | تقرير مميزات SEO المدعومة")
    set_run_font(run, size=9, color="777777")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build()
